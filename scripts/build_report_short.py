# -*- coding: utf-8 -*-
"""
Build a shortened (~35-page main-content) DOCX of the bachelor thesis.

ALL factual content is drawn from the existing report (Report_do_an.pdf) supplied
by the author. No new facts are invented. Where the original had triple-described
material (state machine, RLS/two-client model, cost analysis, design-vs-implementation
duplication), it is stated ONCE. Secondary features are compressed into a single
"Supporting Features" section. Chapters 3-6 of the original are merged into one
"System Design and Implementation" chapter.

Front matter (title, cert, TOC, abbreviations, tables, figures, abstract,
acknowledgements) is preserved per the author's instruction; it does NOT count
toward the 35-page main-content budget.

Figures/tables that were images in the original are inserted as caption +
placeholder note so the author can paste the real image back in the right place.
Pure-text tables (abbreviations, requirements, status, API, results, security,
comparison) are rebuilt as real Word tables.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"c:\Users\ASUS\projects\ai-document-bot\Report_short_35p.docx"

doc = Document()

# ---------- base styles ----------
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(13)
# Slightly tightened from 1.5 -> 1.4 line spacing and 6 -> 4 pt paragraph gap to
# recover whole-document pages without removing any content (layout-only trim).
normal.paragraph_format.line_spacing = 1.4
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
# First-line indent on body paragraphs, matching the original report style
# (~1.27 cm / 0.5 in). Centered/right-aligned lines and captions zero this out.
normal.paragraph_format.first_line_indent = Inches(0.5)

for lvl, sz, sb, sa in [("Heading 1", 16, 12, 8), ("Heading 2", 14, 8, 4), ("Heading 3", 13, 6, 3)]:
    st = doc.styles[lvl]
    st.font.name = "Times New Roman"
    st.font.size = Pt(sz)
    st.font.color.rgb = RGBColor(0, 0, 0)
    st.font.bold = True
    # Trim the generous default heading spacing to recover vertical space.
    st.paragraph_format.space_before = Pt(sb)
    st.paragraph_format.space_after = Pt(sa)
    st.paragraph_format.line_spacing = 1.15
    st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def set_cell_font(cell, bold=False, size=11):
    for p in cell.paragraphs:
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(size)
            r.font.bold = bold
        if not p.runs:
            r = p.add_run("")
            r.font.name = "Times New Roman"
            r.font.size = Pt(size)


def para(text="", style=None, align=None, bold=False, italic=False, size=None, space_after=None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        if size:
            r.font.size = Pt(size)
    if align is not None:
        p.alignment = align
        # centered / right-aligned lines (title page, etc.) should not be indented
        p.paragraph_format.first_line_indent = Inches(0)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text, italic_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if italic_lead:
        r = p.add_run(italic_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def caption(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_after = Pt(6)
    return p


def figure_placeholder(label, desc, width_hint="~75% page width to save space"):
    """A bordered note telling the author to paste the real figure here."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"[ PASTE {label} HERE — {desc} · size: {width_hint} ]")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "bottom", "left", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "dashed")
        e.set(qn("w:sz"), "6")
        e.set(qn("w:space"), "6")
        e.set(qn("w:color"), "AAAAAA")
        pbdr.append(e)
    pPr.append(pbdr)
    return p


def make_table(headers, rows, widths=None, header_size=11, body_size=10):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_font(hdr[i], bold=True, size=header_size)
        # shade header
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "D9D9D9")
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            set_cell_font(cells[i], bold=False, size=body_size)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


def add_toc_field():
    """Insert a real Word TOC field (updates on open / F9)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Right-click and choose “Update Field” to build the table of contents."
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    for el in (fldBegin, instr, fldSep, t, fldEnd):
        run._r.append(el)
    return p


def page_break():
    doc.add_page_break()


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def h3(text):
    doc.add_heading(text, level=3)


# =====================================================================
# FRONT MATTER  (preserved; not part of the 35-page budget)
# =====================================================================

# ----- Title page -----
para("University of Science and Technology of Hanoi", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
para("Department of Information and Communication Technology", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
for _ in range(2):
    para("")
para("BACHELOR THESIS", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=22)
para("")
para("By", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
para("Tran Minh - 23BI14290", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
para("")
para("AI-Assisted Document Approval System with Integrity Verification and Digital Signatures",
     align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
for _ in range(2):
    para("")
para("External Supervisor: Dr. GIANG ANH TUẤN", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
para("Internal Supervisor: Eng. VŨ NGỌC ĐIỆP", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
for _ in range(3):
    para("")
para("Hanoi, July 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
page_break()

# ----- Supervisor certification -----
para("")
para("To whom it may concern,", align=WD_ALIGN_PARAGRAPH.CENTER)
para("")
para("I, ........................................, certify that the thesis / internship report of "
     "Mr. ........................................ is qualified to be presented in the Internship Jury 2025-2026.")
para("")
para("Hanoi, 1/7/2026", align=WD_ALIGN_PARAGRAPH.RIGHT)
para("Supervisor’s signature", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
page_break()

# ----- Table of Contents -----
h1("Table of Contents")
add_toc_field()
page_break()

# ----- List of Abbreviations -----
h1("List of Abbreviations")
abbrevs = [
    ("AI", "Artificial Intelligence"), ("API", "Application Programming Interface"),
    ("DBMS", "Database Management System"), ("FR", "Functional Requirement"),
    ("HTTP", "Hypertext Transfer Protocol"), ("JSON", "JavaScript Object Notation"),
    ("JWT", "JSON Web Token"), ("LLM", "Large Language Model"),
    ("NFR", "Non-functional Requirement"), ("OCR", "Optical Character Recognition"),
    ("PDF", "Portable Document Format"), ("RBAC", "Role-Based Access Control"),
    ("REST", "Representational State Transfer"), ("RLS", "Row-Level Security"),
    ("RSC", "React Server Components"), ("SaaS", "Software as a Service"),
    ("SHA", "Secure Hash Algorithm"), ("ECDSA", "Elliptic Curve Digital Signature Algorithm"),
    ("FIDO", "Fast IDentity Online"), ("TPM", "Trusted Platform Module"),
    ("AAGUID", "Authenticator Attestation Globally Unique Identifier"),
    ("PAdES", "PDF Advanced Electronic Signatures"), ("PKCS", "Public-Key Cryptography Standards"),
    ("SLA", "Service Level Agreement"), ("COSE", "CBOR Object Signing and Encryption"),
    ("SQL", "Structured Query Language"), ("SSR", "Server-Side Rendering"),
    ("UI", "User Interface"), ("UML", "Unified Modeling Language"),
    ("URL", "Uniform Resource Locator"), ("UX", "User Experience"),
]
make_table(["Abbreviation", "Meaning"], abbrevs, widths=[1.6, 4.4])
page_break()

# ----- List of Tables -----
h1("List of Tables")
for line in [
    "Table 2.1  Feature comparison of representative document approval and AI assistance systems.",
    "Table 3.1  Summary of the main database tables in the system.",
    "Table 3.2  Summary of functional requirements.",
    "Table 3.3  Summary of non-functional requirements.",
    "Table 3.4  Document status values and transitions.",
    "Table 3.5  Main REST API endpoints exposed by the backend.",
    "Table 4.1  Coverage of functional requirements in the implemented prototype.",
    "Table 4.2  Measured response times for representative interactions under light load.",
    "Table 4.3  Security attack scenarios tested against the prototype.",
]:
    para(line, size=12)
page_break()

# ----- List of Figures -----
h1("List of Figures")
for line in [
    "Figure 3.1  Document state machine showing transitions between status values.",
    "Figure 3.2  Use case diagram of the AI-assisted document approval system.",
    "Figure 3.3  Entity-relationship diagram of the database schema.",
    "Figure 3.4  Overall three-tier system architecture.",
    "Figure 3.5  Sequence diagram of the document submission and review flow.",
    "Figure 3.6  Hybrid text extraction pipeline with OCR fallback.",
    "Figure 3.7  Staging-then-validate upload pipeline.",
    "Figure 3.8  Verification panel across three tamper scenarios in the multi-party WebAuthn signing flow.",
    "Figure 4.1  Login and registration interface.",
    "Figure 4.2  Employee dashboard with document overview.",
    "Figure 4.3  Document detail page with inline PDF viewer.",
    "Figure 4.4  AI assistant panel showing the generated summary.",
    "Figure 4.5  Approval certificate and integrity verification page.",
]:
    para(line, size=12)
page_break()

# ----- Abstract (kept verbatim - already <=250 words) -----
h1("Abstract")
para(
    "This report presents the design and implementation of an AI-assisted document approval "
    "system that combines a multi-role review workflow, integrity verification, and digital "
    "signatures. It addresses the weaknesses of traditional review processes, in which email- and "
    "paper-based workflows are hard to track, lack accountability, and offer no automated help in "
    "analyzing document content. The system is built as a modern web application using Next.js 16, "
    "React 19, and Tailwind CSS on the client, with Supabase (PostgreSQL with Row-Level Security, "
    "Auth, and Storage) and the OpenAI API on the server. It supports three roles — employee, "
    "reviewer, and administrator — and enforces access control at both the application and database "
    "layers. Documents pass through a multi-round, multi-reviewer approval pipeline; once approved, "
    "the system records multi-party WebAuthn digital signatures, one from the owner at submission and "
    "one from each approving reviewer, produced by hardware-bound private keys in the device TPM, "
    "together with a printable certificate. The integrated OpenAI-powered assistant extracts text from "
    "uploaded PDFs, generates structured summaries with key points and risk notes, and answers grounded "
    "questions about document content. Experimental evaluation on a representative set of documents shows "
    "that the system handles the complete review lifecycle correctly, generates AI summaries within a few "
    "seconds, and prevents unauthorized data access through layered security controls."
)
p = doc.add_paragraph()
r = p.add_run("Keywords: ")
r.bold = True
p.add_run("document approval, digital signatures, integrity verification, role-based access control, "
          "WebAuthn, large language models.")
page_break()

# ----- Acknowledgements -----
h1("Acknowledgements")
para(
    "I would like to express my sincere gratitude to my supervisors, Mr. Giang Anh Tuan and Mr. Vu "
    "Ngoc Diep, for their continuous guidance, valuable feedback, and encouragement throughout the "
    "development of this graduation project. Their suggestions on system architecture, data modeling, "
    "security design, and the integration of large language models greatly improved the quality of this "
    "work and helped me overcome several technical and conceptual challenges encountered along the way."
)
para(
    "I also thank the lecturers and staff of the University of Science and Technology of Hanoi, in "
    "particular the Department of Information and Communication Technology and the ICT Lab, for "
    "providing the academic environment, technical infrastructure, and support that made this project "
    "possible. The knowledge gained from courses on databases, software engineering, web technologies, "
    "and security formed the foundation on which this project was built."
)
para(
    "Finally, I am deeply grateful to my family, classmates, and friends for their patience, "
    "understanding, and motivation during the time I dedicated to designing, implementing, and testing "
    "this system."
)
page_break()

# =====================================================================
# CHAPTER 1 - INTRODUCTION   (target ~5 pages)
# =====================================================================
h1("Chapter 1  Introduction")

h2("1.1  Background and Motivation")
para(
    "In modern organizations, the review and approval of documents such as reports, contracts, internal "
    "proposals, and policy drafts is a fundamental administrative process. Documents flow from an author "
    "to one or more reviewers, who read the content, provide feedback, and either approve the document or "
    "request revisions. In many workplaces this process is still carried out manually through email "
    "attachments, printed copies, and paper-based signatures, which leads to several recurring problems: "
    "documents are difficult to track, feedback is scattered across email threads, revision history is "
    "hard to reconstruct, and there is no reliable way to verify after the fact that a signed document "
    "has not been altered."
)
para(
    "These challenges are compounded by the increasing volume and complexity of documents. Reviewers "
    "often must read long PDFs with technical or legal content under tight deadlines, while traditional "
    "tools provide little automated support: a PDF reader can display a file but cannot summarize it, "
    "highlight risky clauses, or answer questions about it. Recent advances in large language models "
    "(LLMs) accessible through cloud APIs can summarize long documents, extract structured information, "
    "and answer grounded questions, and when integrated carefully into a review workflow they can reduce "
    "the cognitive load on reviewers without replacing their final decision-making authority. At the same "
    "time, modern web platforms such as Next.js and managed services such as Supabase make it feasible "
    "for a single student to build a production-quality application that combines secure authentication, "
    "fine-grained access control, file storage, and AI integration within a reasonable timeframe."
)

h2("1.2  Problem Statement and Existing-System Gap")
para(
    "Despite the digitization of office work, many organizations still rely on ad hoc combinations of "
    "email, shared folders, and instant messaging for document review. Such practices suffer from a lack "
    "of workflow visibility (no single place to see who approved or rejected and when), weak "
    "accountability (approvals recorded as informal replies that can be modified, lost, or forged), no "
    "integrity guarantee (no automated way to verify later that a file is exactly the approved version), "
    "no intelligent assistance (reviewers must read entire documents with no summary or question "
    "answering), and unstructured multi-reviewer decisions (it is unclear whether all reviewers must "
    "agree or how a single rejection ends a round)."
)
para(
    "Existing general-purpose tools each address only part of the problem and treat the rest as omissions "
    "or paid integration points. Electronic-signature platforms sign PDFs but offer no AI; document "
    "management systems handle workflow but bolt signing and AI on; AI document assistants summarize but "
    "have no concept of review or signing; and certificate authorities provide legally-compliant signing "
    "but no workflow or AI. This separation is the gap the project targets: there is no single, "
    "low-overhead application that unifies workflow management, AI assistance, and integrity "
    "verification. Section 2.5 makes this comparison concrete in Table 2.1."
)

h2("1.3  Objectives")
para("The primary objectives of this project are as follows:")
bullet("Design and implement a web-based document approval system supporting three roles — employee, "
       "reviewer, and administrator — each with appropriate permissions and dashboards.")
bullet("Model the document lifecycle as a clearly defined state machine (draft, pending, approved, "
       "rejected) with transitions implemented through dedicated server-side endpoints.")
bullet("Implement a multi-round, multi-reviewer approval pipeline in which submission opens a new round, "
       "all assigned reviewers must approve for the document to be approved, and any single rejection "
       "immediately terminates the round.")
bullet("Integrate the OpenAI API for summarization, key-point and risk-note extraction, and grounded "
       "question answering on uploaded PDFs, with cost controls and a fallback path for scanned documents.")
bullet("Implement a multi-party digital signing system in which the owner signs at submission and each "
       "reviewer signs at approval, using WebAuthn / FIDO2 platform authenticators so that each signature "
       "is bound to the signer’s device via a non-extractable TPM-backed private key, with a printable "
       "certificate and a one-click verify button.")
bullet("Enforce access control both at the application level (page guards, role-aware UI) and at the "
       "database level (Row-Level Security in PostgreSQL), so that the system remains secure even with a "
       "compromised client.")

h2("1.4  Scope and Limitations")
para(
    "The project focuses on a single-organization document review platform delivered as a web "
    "application, managing PDF documents only, with a structured workflow, an AI assistant powered by an "
    "external LLM, and integrity certificates based on cryptographic hashing and WebAuthn signatures. "
    "The intended deployment is a small to medium organization or a university department."
)
para(
    "Several limitations apply. The system supports PDF only; other formats are out of scope. Digital "
    "signing relies on platform authenticators (Windows Hello, Touch ID, mobile biometric), so a Linux "
    "desktop without one must use another device. PDF-embedded signatures (PAdES / PKCS#7) are out of "
    "scope; signatures live in a separate table and are verified via the certificate page. The OCR "
    "fallback for scanned PDFs is limited to ten pages to keep API costs predictable. The system depends "
    "on the OpenAI API, whose availability and pricing are outside the project’s control. Evaluation "
    "focuses on functional correctness and qualitative usability rather than large-scale user studies."
)
para(
    "The system has a usage-based cost component. Hosting fits the free tiers of Vercel and Supabase at "
    "this scale, but every AI call incurs a per-call OpenAI charge. Based on per-call token usage logged "
    "automatically in the audit_logs table (not estimates), a structured summary costs approximately "
    "0.33 US cent, a grounded question approximately 0.22 US cent, and an OCR extraction approximately "
    "1.41 US cent; sample sizes are modest (N = 7–12 per operation) and per-call cost is bounded by the "
    "12,000-character input truncation and the ten-page OCR cap. For a small organization processing 100 "
    "documents per month (one summary and five questions each, about a quarter requiring OCR), the "
    "resulting monthly AI cost is under 2 US dollars — modest relative to the per-user subscriptions of "
    "comparable commercial products. Document text is sent to OpenAI only on explicit user action and is "
    "governed by the provider’s API terms (excluded from training, retained at most thirty days); "
    "organizations with stricter data-sovereignty requirements, for example under Vietnam’s Decree "
    "13/2023/ND-CP, would substitute a self-hosted model, identified as future work in Chapter 5. A "
    "per-user rate limit on the AI endpoints (Section 3.6) further bounds request volume."
)

h2("1.5  Research Questions")
para("Based on the identified problem and objectives, the project addresses three research questions:")
bullet("How can a multi-role document approval workflow be modeled as a state machine and implemented in "
       "a modern web stack such that it is both expressive enough for realistic organizational use and "
       "simple enough for a single developer to maintain?", italic_lead="RQ1: ")
bullet("How can large language model capabilities be integrated into a document review interface in a way "
       "that is genuinely useful for reviewers, controls API costs, and degrades gracefully when the "
       "model fails or the input is not machine-readable?", italic_lead="RQ2: ")
bullet("To what extent can layered security controls, combining application-level role checks with "
       "database-level Row-Level Security, prevent unauthorized access to documents and approval records, "
       "including when the client is malicious or compromised?", italic_lead="RQ3: ")

h2("1.6  Methodology Overview")
para(
    "The project follows an end-to-end methodology combining requirements analysis, system design, "
    "implementation, and evaluation. It begins with an analysis of typical document review scenarios in "
    "small organizations, which informs the identification of actors, use cases, and functional and "
    "non-functional requirements. The system is then designed as a three-tier web application with a "
    "clear separation between the user interface, the application logic, and the data layer. The data "
    "layer is a managed PostgreSQL database governed by Row-Level Security; the application logic is "
    "implemented as server components and API routes in Next.js; and the user interface is built with "
    "React and Tailwind CSS using a server-component-first approach. The system is evaluated functionally "
    "by walking through each use case with multiple role accounts, the AI integration qualitatively and "
    "by measuring response times, and security by attempting unauthorized access both through the UI and "
    "by using the anonymous database client directly."
)
page_break()

# =====================================================================
# CHAPTER 2 - BACKGROUND & RELATED WORK   (target ~6 pages)
# =====================================================================
h1("Chapter 2  Background and Related Work")

h2("2.1  Document Approval Workflows as State Machines")
para(
    "Document management and approval systems have been of practical and academic interest since the "
    "early days of office automation. Traditional document management systems focus on storage, "
    "versioning, and basic access control, while workflow management systems model the document lifecycle "
    "as a sequence of activities with transitions triggered by user actions. A common abstraction is to "
    "represent a document’s lifecycle as a state machine, where each state corresponds to a stage in the "
    "review process and transitions correspond to user-initiated events such as submission or decision. "
    "This abstraction makes the workflow explicit and visible, lets the system enforce invariants such as "
    "preventing modification of an approved document, and provides a natural place to attach side effects "
    "such as notifications and audit log entries. The system in this project occupies an intermediate "
    "position between oversimplified single-decision tools and highly configurable enterprise systems: "
    "the workflow is fixed and intentionally simple, with multi-reviewer rounds and a clear set of states "
    "sufficient for the common case of internal organizational review."
)

h2("2.2  Role-Based Access Control and Row-Level Security")
para(
    "Role-Based Access Control (RBAC) is one of the most widely used authorization models in enterprise "
    "software: permissions are assigned to roles rather than directly to users, and users acquire "
    "permissions through role membership, simplifying administration. In web applications, RBAC is "
    "traditionally enforced at the application layer by checking the user’s role on every request. This "
    "is effective but has an important weakness: any bug in the application code can bypass the control "
    "and expose data. A defense-in-depth strategy adds enforcement at the database layer through "
    "Row-Level Security (RLS), a feature that lets the administrator define, in SQL, predicates that "
    "determine which rows a given user can see or modify. When a query is executed, the database engine "
    "implicitly augments it with these predicates, so even a permissive query such as "
    "SELECT * FROM documents returns only rows that satisfy the policy. PostgreSQL supports RLS natively, "
    "and this project uses both layers: every route checks the authenticated user’s role, and the "
    "database enforces RLS policies on the documents, document_versions, document_highlights, and related "
    "tables. If an attacker bypasses the application logic, the database still refuses to return "
    "unauthorized rows. The implementation of these policies is described in Section 3.2."
)

h2("2.3  Cryptographic Hashing and Hardware-Bound Digital Signatures")
para(
    "Cryptographic hash functions such as SHA-256 are a foundational tool for verifying file integrity: "
    "the output is deterministic and any change to the input produces a completely different output with "
    "overwhelming probability. In a document approval setting, the system can compute and store the "
    "SHA-256 hash of an approved file; anyone holding the file can later recompute the hash and compare "
    "it with the stored value to detect tampering. However, a hash alone proves only that a file is "
    "unchanged given that one trusts the system storing the hash; it does not prove who approved the "
    "document or that the approver was not impersonated. Full non-repudiation requires a public-key "
    "digital signature in which the approver signs with a private key only they possess."
)
para(
    "This project implements non-repudiation through the WebAuthn (FIDO2) standard, which defines a "
    "browser API and a server-side verification protocol for public-key authentication using "
    "hardware-backed authenticators. Keys are generated and stored in the device’s Trusted Platform "
    "Module (TPM) on Windows, the Secure Enclave on Apple devices, or equivalent secure hardware on "
    "mobile platforms. The private key is non-extractable by design, and every signing operation requires "
    "fresh user-verification via PIN or biometric. The system uses WebAuthn assertions (ECDSA P-256 over "
    "SHA-256) as the signing primitive, with the file hash serving as the challenge so that each "
    "signature cryptographically binds a specific signer to specific document content. The signature "
    "counter in WebAuthn’s authenticatorData provides additional defense against cloned-authenticator "
    "replay attacks."
)

h2("2.4  Large Language Models for Document Understanding")
para(
    "Large language models, particularly those of the transformer family, achieve strong results on "
    "summarization, question answering, and information extraction, and are accessible through cloud APIs "
    "that accept a prompt and return a generated response. The model used in this project is "
    "gpt-5.4-mini, accessed through the official OpenAI SDK for JavaScript and TypeScript; the choice of a "
    "small, fast model reflects the cost-control requirements of Chapter 1, and the model identifier is "
    "read from an environment variable so it can be replaced without code changes. Integrating an LLM "
    "into a document review system poses three practical challenges: documents are PDFs that may contain "
    "a text layer, scanned images, or both, so a robust pipeline must extract text and fall back to OCR; "
    "models charge per token, so input must be bounded to control cost; and model outputs are "
    "non-deterministic, so a JSON-schema-constrained structured-output mode is used to make responses "
    "reliably parseable. The application takes advantage of all three techniques, as described in "
    "Section 3.4. A trade-off inherent to this design is that document content is transmitted to an "
    "external service when AI features are invoked; the provider excludes API data from training and "
    "imposes a thirty-day retention limit, but a self-hosted model would be the appropriate substitute "
    "for highly sensitive documents."
)

h2("2.5  Comparison with Existing Systems")
para(
    "Several commercial products address parts of the same problem space, but each focuses on one of the "
    "three concerns — workflow, AI assistance, or digital signing — and treats the other two as "
    "integration points or omissions. Table 2.1 groups representative systems by their primary focus and "
    "contrasts them with this project."
)
comp_headers = ["System / Category", "Approval workflow", "AI summarization & Q&A", "Signing type", "Deployment & cost"]
comp_rows = [
    ["This project",
     "Multi-round, multi-reviewer with deadlines and resubmission",
     "Real LLM (gpt-5.4-mini) with structured outputs; hybrid OCR fallback",
     "Multi-party WebAuthn / FIDO2 (TPM-backed ECDSA P-256)",
     "Free-tier hosting (Vercel + Supabase); pay-per-use AI (~0.3¢ summary, ~0.2¢ question, ~1.4¢ OCR)"],
    ["DocuSign, Adobe Acrobat Sign, PandaDoc",
     "Built-in routing and reminders",
     "None (or limited add-ons)",
     "Click-to-sign with audit trail; eIDAS / ESIGN compliant",
     "Commercial SaaS, ~$10–25 / user / month"],
    ["M-Files, DocuWare, SharePoint Approvals",
     "Configurable workflows, document lifecycle management",
     "None natively (Copilot as paid add-on)",
     "Bolted-on e-signature integrations",
     "Commercial license; enterprise-oriented"],
    ["Humata.ai, ChatPDF, Acrobat AI Assistant",
     "None — read-only AI on uploaded files",
     "Native — summarization and grounded Q&A",
     "None",
     "Freemium SaaS"],
    ["VNPT-CA, Viettel-CA, BKAV ChukyDienTu",
     "None — pure PKI signing service",
     "None",
     "PKI certificate (PAdES, X.509) issued by CA",
     "~500,000–1,000,000 VND / user / year"],
]
make_table(comp_headers, comp_rows, widths=[1.3, 1.2, 1.3, 1.3, 1.4], body_size=9)
caption("Table 2.1: Feature comparison of representative document approval and AI assistance systems.")
para(
    "Dedicated electronic-signature platforms produce legally-binding PAdES signatures but offer little "
    "native AI and only simple routing. Enterprise document management systems handle complex workflows "
    "but require months of configuration and bolt signing and AI on. Pure AI assistants summarize "
    "effectively but have no concept of submission, review, or signing. Vietnamese certificate "
    "authorities deliver legally-compliant signing but no workflow or AI. The proposed system "
    "intentionally co-locates the three concerns — multi-round approval, AI-assisted reading, and "
    "hardware-bound signing — in a single application deployable on free-tier infrastructure by a single "
    "developer. The trade-off is that the workflow is fixed rather than configurable, and the signatures "
    "are not yet embedded in the PDF binary; both are noted as future work in Chapter 5."
)
page_break()

# =====================================================================
# CHAPTER 3 - SYSTEM DESIGN & IMPLEMENTATION  (merged 3+4+5+6; target ~16 pages)
# =====================================================================
h1("Chapter 3  System Design and Implementation")

h2("3.1  Technology Stack and Architecture")
para(
    "The implementation rests on three frameworks — Next.js 16, React 19, and Tailwind CSS v4 — and a "
    "small set of supporting libraries: the Supabase JavaScript client, the OpenAI SDK, "
    "@simplewebauthn/server, pdf-parse, and react-pdf. Next.js 16 with the App Router is the principal "
    "framework, responsible for file-system routing, the server-component runtime, and the API route "
    "handlers. It was chosen over a manually composed Express + React stack because the App Router "
    "encourages a server-component-first design that moves data fetching to the server and keeps "
    "sensitive logic out of the client, integrates cleanly with Vercel for deployment with "
    "zero-configuration HTTPS, and supports TypeScript route handlers so REST-like endpoints map directly "
    "to files in the source tree. React 19 provides the user interface; styling uses Tailwind CSS v4 with "
    "a small custom design system (page-shell layout, section cards, metric cards, status pills, button "
    "variants) rather than a generic component library, keeping the bundle small. PDFs are rendered "
    "inline by react-pdf (wrapping pdfjs-dist). The entire application is written in strict-mode "
    "TypeScript, which catches null and undefined errors at compile time — particularly valuable in the "
    "WebAuthn verification path, where a missing check could expose the system to forged signatures."
)
para(
    "The system follows a three-tier architecture composed of a browser-based front end, a Next.js "
    "server, and a Supabase data backend. The front end and server are deployed together as a single "
    "Next.js application on Vercel, while the database, authentication, and storage services are provided "
    "by Supabase. The AI assistant is implemented by server-side code that calls the OpenAI API. Figure "
    "3.4 summarizes the architecture, and Table 3.1 (below) and the component breakdown describe the "
    "layers."
)
figure_placeholder("Figure 3.4", "Overall three-tier system architecture (was Figure 4.4)")
caption("Figure 3.4: Overall three-tier system architecture.")
para(
    "A defining feature of the architecture is the use of two distinct database clients on the server "
    "side. The anon-key client respects Row-Level Security and is used for all read operations performed "
    "by server components when rendering pages: because the user’s identity is encoded in the JSON Web "
    "Token accompanying the request, the database automatically restricts the result set to rows the user "
    "is entitled to see. The service-role client bypasses RLS and is used only inside dedicated API "
    "routes for privileged writes such as inserting documents or notifications, always after the route "
    "has verified the caller’s identity and permissions. So reads are governed declaratively by the "
    "database and writes imperatively by server code, and the browser can bypass neither. In addition to "
    "the request-response path, the architecture includes a websocket channel between the browser and "
    "Supabase Realtime for live notification and dashboard updates."
)
para(
    "Authentication is handled by Supabase Auth, which stores credentials, hashes passwords with a salted "
    "scheme, and issues JSON Web Tokens as HTTP-only cookies inaccessible to client-side JavaScript and "
    "therefore robust against cross-site scripting. A helper module exposes two functions used throughout "
    "the server-side code: requireUser, which fetches the current profile and redirects unauthenticated "
    "or unapproved users, and requireRole, which additionally checks the profile’s role. These helpers "
    "are called at the top of every protected server component and the start of every API route, so no "
    "protected page or endpoint executes without a verified identity. Combined with RLS, this creates a "
    "layered access control system in which the database refuses unauthorized reads even if a defect in "
    "the application code allows an unauthorized request to reach it."
)

h2("3.2  Data Model and Storage Design")
para(
    "All structured data is stored in the managed PostgreSQL database, while binary PDF files are stored "
    "in Supabase Storage; both services share the same access tokens and client libraries. The schema "
    "consists of ten tables: nine model the core concepts of the document approval domain, and a tenth, "
    "document_ai_messages, is an auxiliary log of AI question-and-answer interactions. The tables are "
    "normalized to third normal form, with foreign keys and carefully chosen cascade rules. Table 3.1 "
    "summarizes them."
)
db_headers = ["Table", "Role in the system"]
db_rows = [
    ["profiles", "Per-user information extending the built-in auth table: full name, role (employee, "
     "reviewer, admin), account status, and — once Windows Hello signing is set up — the WebAuthn "
     "credential ID, COSE public key, sign counter, AAGUID, device-type, and registration timestamp."],
    ["documents", "A single document with title, description, owner, current status, and timestamps. The "
     "approved_hash column stores the SHA-256 hash computed at the moment of unanimous approval, used to "
     "detect post-approval tampering."],
    ["document_versions", "Every uploaded version: storage path, extracted text, a cached SHA-256 content "
     "hash, and the uploader."],
    ["approvals", "Each assigned reviewer’s decision in each round: round number, status (pending / "
     "approved / rejected), optional comment, and deadline."],
    ["document_highlights", "Passage-level comments on a specific version: page number, selected text, "
     "comment body, and bounding rectangles for rendering."],
    ["document_signatures", "One row per signing event: signer, role (owner_submission or "
     "reviewer_approval), round number, SHA-256 hash of the signed file, WebAuthn signature bytes, "
     "authenticatorData, and clientDataJSON — used by the verify endpoint and the certificate page."],
    ["document_ai_results", "Caches AI output: extracted text, summary, key points, risk notes, and usage "
     "metadata."],
    ["document_ai_messages", "Q&A history of the AI assistant — one row per question (document, asking "
     "user, question text, model answer)."],
    ["notifications", "In-app notifications: type, title, body, optional related document, and a "
     "read/unread flag."],
    ["audit_logs", "Significant actions: actor, action code, optional target reference, JSONB metadata, "
     "and timestamp; used by the administrator panel."],
]
make_table(db_headers, db_rows, widths=[1.5, 4.5], body_size=9)
caption("Table 3.1: Summary of the main database tables in the system.")
para(
    "Each document has exactly one owner (documents.owner_id → profiles.id, non-null and immutable, "
    "reflecting that authorship does not transfer) and accumulates one or more document_signatures rows: "
    "one for the owner at submission and one for each approving reviewer. A partial unique index on "
    "(document_id, signer_id, signature_role, round_no) guarantees a party cannot sign the same document "
    "twice in the same role within one round, while a composite unique constraint on (document_id, "
    "reviewer_id, round_no) prevents a reviewer being assigned twice in the same round. Highlights are "
    "version-scoped so they remain attached to the version they were authored on. Notifications and audit "
    "logs are deliberately decoupled, each carrying only a nullable document reference. The "
    "entity-relationship diagram is shown in Figure 3.3."
)
figure_placeholder("Figure 3.3", "Entity-relationship diagram of the database schema (was Figure 4.3)")
caption("Figure 3.3: Entity-relationship diagram of the database schema.")
para(
    "Row-Level Security policies are an integral, version-controlled part of the schema. For the "
    "documents table, the SELECT policy makes a row visible only if the calling user is the owner, has at "
    "least one approval assignment for the document, or is an administrator; document_versions inherits "
    "visibility from the parent document via an EXISTS subquery. For approvals, SELECT is visible to the "
    "document owner, the assigned reviewer, and administrators. Notifications are restricted so each user "
    "sees only their own; audit logs are restricted to administrators. Inserts and updates are generally "
    "not permitted through the anon client — they go through API routes running with the service-role key "
    "after explicit authorization checks. This yields a clear model: read access governed declaratively "
    "by the database, write access governed imperatively by server code."
)
para(
    "Binary PDFs are stored in a private Supabase Storage bucket with paths organized hierarchically as "
    "user_id / document_id / timestamp-name, which makes user-scoped storage policies easy to express and "
    "lets file deletion remove a single prefix. A staging area under each user prefix supports a "
    "validate-then-commit upload pipeline, described in Section 3.3."
)

h2("3.3  Upload Pipeline")
para(
    "Both creating a new document and adding a new version use the same staging-then-validate pattern, "
    "shown in Figure 3.7. The client performs simple pre-flight checks (PDF extension, size within the "
    "configured maximum) and uploads the file to a staging path under the user’s storage prefix, then "
    "sends a request to a server endpoint with the staging path and document metadata. On the server, the "
    "validation helper fetches only the first few bytes of the staged file via an HTTP Range request "
    "(avoiding a full download), checks that the bytes begin with the PDF magic number %PDF-, and "
    "confirms the total size is within the limit. If validation fails, the staging object is deleted and "
    "the endpoint returns HTTP 400. If it succeeds, the endpoint uses the service-role client to move the "
    "file to its final path and insert the corresponding row in the documents or document_versions table."
)
figure_placeholder("Figure 3.7", "Staging-then-validate upload pipeline (was Figure 6.1)")
caption("Figure 3.7: Staging-then-validate upload pipeline.")
para(
    "The crucial property of this design is that the insert into documents and document_versions happens "
    "only on the server, through the service-role client, after validation; the anon-key client in the "
    "browser cannot insert directly into these tables. Even a maliciously crafted client that uploads a "
    "file with a fake extension or embedded executable content cannot create a database row pointing to "
    "it: the file must first survive server-side validation."
)

h2("3.4  Workflow: State Machine and Multi-Reviewer Pipeline")
para(
    "The lifecycle of a document is modeled as a finite state machine with four states, enforced both by "
    "the database (the status column is constrained to these values) and by the API routes (which check "
    "the current state before applying a transition). Table 3.4 lists the states and the events that "
    "cause transitions, and Figure 3.1 illustrates them."
)
st_headers = ["Status", "Triggered by", "Meaning"]
st_rows = [
    ["draft", "Initial creation, or upload after rejection",
     "Created but not yet submitted. Only the owner can edit and upload new versions."],
    ["pending", "Owner signs and submits for review",
     "One or more reviewers assigned. The owner’s WebAuthn signature has been recorded "
     "(signature_role = owner_submission). At least one reviewer has not yet decided."],
    ["approved", "All reviewers approve and sign",
     "Every assigned reviewer has approved the latest version and produced a WebAuthn signature "
     "(reviewer_approval). An approval-time hash (approved_hash) is also recorded. Terminal state."],
    ["rejected", "Any reviewer rejects",
     "One reviewer has rejected, terminating the round immediately. The owner can upload a new version, "
     "which returns the document to draft."],
]
make_table(st_headers, st_rows, widths=[1.0, 1.7, 3.3], body_size=10)
caption("Table 3.4: Document status values and transitions.")
figure_placeholder("Figure 3.1", "Document state machine (was Figure 4.1)")
caption("Figure 3.1: Document state machine showing transitions between status values.")
para(
    "Many approval scenarios involve more than one reviewer per document. The system supports this "
    "through a round-based design: when the owner submits, they select one or more reviewers, and the "
    "system creates one approvals row per reviewer with the same round number (one plus the current "
    "maximum) and transitions the document to pending. Each reviewer independently approves or rejects "
    "through a dedicated route; after each decision the route recomputes the aggregate status. If any "
    "reviewer rejects, the document becomes rejected and remaining pending approvals are superseded; if "
    "all approve, it becomes approved and the approval-time hash is stored. Resubmission after rejection "
    "is a natural continuation of the round model: the owner uploads a new version (resetting the "
    "document to draft) and submits again with a possibly different set of reviewers, incrementing the "
    "round number so history is preserved and displayed grouped by round."
)
para(
    "Separation of duties is a deliberate design point: approval authority comes not from a job title but "
    "from being assigned as a reviewer on that specific document, and the submission route blocks the "
    "owner from assigning themselves, so no one can approve their own document. The submission route is "
    "located at /api/documents/[id]/submit and the decision route at "
    "/api/approvals/[approvalId]/decide; both verify the caller’s authorization, the current state, and "
    "round membership before writing, and append a row to audit_logs. Figure 3.5 shows the sequence of a "
    "document submission and a reviewer decision, including the signature step described in Section 3.6."
)
figure_placeholder("Figure 3.5", "Sequence diagram of submission and review flow (was Figure 4.5)")
caption("Figure 3.5: Sequence diagram of the document submission and review flow.")

h2("3.5  AI Assistant: Hybrid Extraction, Summarization, and Q&A")
para(
    "Reliable text extraction is a prerequisite for the AI features. Uploaded PDFs vary: most contain a "
    "proper text layer that can be parsed directly, but some are scanned images with no text. The system "
    "implements a hybrid extraction pipeline, shown in Figure 3.6. The file is read into memory and "
    "passed to the pdf-parse library, which returns the concatenated text of all pages. If the output "
    "exceeds a threshold of one hundred characters, the pipeline treats it as valid; otherwise it "
    "concludes the PDF is likely a scan and switches to the OCR fallback, uploading the original buffer "
    "to the OpenAI Files API for the multimodal model to read. To keep API costs predictable, the OCR "
    "fallback refuses PDFs with more than ten pages and returns HTTP 413; text-based PDFs of any length "
    "are unaffected. Each call records which path was used in audit_logs, useful both for cost analysis "
    "and for understanding where summary quality may be lower due to imperfect OCR."
)
figure_placeholder("Figure 3.6", "Hybrid text extraction pipeline with OCR fallback (was Figure 5.1)")
caption("Figure 3.6: Hybrid text extraction pipeline with OCR fallback.")
para(
    "Once the text is extracted, the AI workspace offers two operations through dedicated API routes that "
    "call the OpenAI client through a small wrapper module. The summary endpoint "
    "(/api/documents/[id]/ai-summary) sends the extracted text, truncated to a configurable maximum of "
    "twelve thousand characters, with a prompt instructing the model to return three fields — a concise "
    "summary, a list of key points, and a list of risk notes — constrained by a JSON schema so the model "
    "returns a valid object with exactly those fields. The result is stored as a row in "
    "document_ai_results, which acts as a cache: subsequent requests load the most recent row instead of "
    "regenerating. The question endpoint (/api/documents/[id]/ai-chat) takes a natural-language question, "
    "truncates it to one thousand characters, and constructs a prompt that includes the extracted text "
    "with explicit instructions to ground all answers in that text; each question and the model’s answer "
    "are recorded in document_ai_messages, giving the document a persistent Q&A history. Both endpoints "
    "log the model name and token usage to audit_logs. The default model is gpt-5.4-mini, with a separate "
    "optional OPENAI_OCR_MODEL variable for the OCR path. The AI is assistance, not authority — the human "
    "reviewer always makes the final decision."
)
para(
    "Errors from the OpenAI API are mapped to appropriate HTTP status codes by the wrapper: 503 when the "
    "API key is not configured, 429 when the rate limit is reached, 402 when the project quota is "
    "exhausted, and 413 when an OCR fallback is requested for a document over the page limit. These are "
    "translated into user-friendly messages in the front end."
)

h2("3.6  Digital Signature and Integrity Verification")
para(
    "The system implements a multi-party digital signing model based on WebAuthn / FIDO2, in which each "
    "participating party produces a cryptographic signature with a hardware-bound private key in the "
    "device’s TPM. Signing is woven directly into the workflow rather than treated as a separate terminal "
    "action: the owner signs at submission to attest that the uploaded file is the version they intended "
    "to circulate, and each approving reviewer signs at the moment of approval to attest that the file "
    "they reviewed is the one they endorse. A document is fully signed only when the owner’s submission "
    "signature and every assigned reviewer’s approval signature have been recorded; rejections do not "
    "require a signature, since rejection terminates rather than commits the round."
)
para(
    "The first time a user signs, the application opens a Windows Hello (or platform-equivalent biometric "
    "/ PIN) prompt and generates an ECDSA P-256 keypair inside the device’s TPM, using "
    "navigator.credentials.create() with authenticatorAttachment: \"platform\" and userVerification: "
    "\"required\". The server stores the user’s COSE-encoded public key, AAGUID, sign counter, "
    "device-type indicator, and registration timestamp in profiles. The private key is non-extractable: "
    "it never leaves the TPM and cannot be exported even by the user. Each subsequent signing event "
    "proceeds in three steps: the client requests the SHA-256 hash of the latest file version from a "
    "dedicated endpoint; it calls navigator.credentials.get() with the file hash as the WebAuthn "
    "challenge, triggering Windows Hello so the TPM signs and returns the signature bytes, "
    "authenticatorData, and clientDataJSON; and it posts all three values together with the action "
    "payload to the corresponding server route."
)
para(
    "The server route (/api/documents/[id]/submit for owner signatures, "
    "/api/approvals/[approvalId]/decide for reviewer signatures) verifies the assertion before writing "
    "any database row. It imports the signer’s public key, reconstructs the file hash, and calls "
    "verifyAuthenticationResponse from @simplewebauthn/server, which checks that the challenge inside "
    "clientDataJSON matches the file hash, that the origin and relying-party identifier match the "
    "deployed application, that the user-verification flag is set, and that the signature is "
    "mathematically valid under the signer’s public key. The route then updates the stored sign counter "
    "(to detect cloned-authenticator replays) and inserts the signature row. Only on a positive "
    "verification does any state transition occur; a malformed or forged signature causes the entire "
    "submission or approval to be rejected with no side effects."
)
para(
    "Two user-facing features expose the signature evidence. The Verify Integrity button on the document "
    "detail page recomputes the SHA-256 hash of the current file in storage, fetches every signature, and "
    "runs verifyAuthenticationResponse against each signer’s stored public key with the signature’s "
    "recorded hash as the expected challenge. The interface presents two distinct judgements per "
    "signature: a Hash Match / Mismatch badge proving whether the current file is identical to what each "
    "party signed, and a WebAuthn-ES256 Valid / Invalid badge proving whether the signature was produced "
    "by the holder of that signer’s registered private key. The certificate page displays the same "
    "verification augmented with attestation evidence decoded from authenticatorData (a friendly "
    "authenticator name from the AAGUID, the User-Verified flag, the device-type indicator, and the sign "
    "counter). Figure 3.8 illustrates the three observable outcomes."
)
figure_placeholder("Figure 3.8", "Verification panel across three tamper scenarios (was Figure 5.2, panels a/b/c)")
caption("Figure 3.8: Verification panel across three tamper scenarios in the multi-party WebAuthn signing "
        "flow. (a) No tampering — every signature shows Hash Match and WebAuthn-ES256 Valid. (b) File "
        "modified after submission but before approval — only the owner’s signature exists; it shows Hash "
        "Mismatch but remains Valid. (c) File modified after final approval — both signatures show Hash "
        "Mismatch but remain Valid.")
para(
    "The combination of the two judgements is diagnostic. A valid-but-mismatched signature tells the "
    "viewer precisely that nobody was impersonated but the file was altered after it was signed — a "
    "distinction a single hash check could never make. The design provides three jointly-enforced "
    "security properties. Authenticity: a signature can only be produced by a user who physically "
    "possesses the device holding the private key and passes user-verification at signing time. "
    "Integrity: the file hash is incorporated into the signature via the challenge, so any modification "
    "after signing is detected by hash mismatch. Non-repudiation: the public key is registered to a "
    "specific account during a Windows Hello ceremony that itself required biometric or PIN, so a signer "
    "cannot plausibly deny producing a signature that verifies under their public key. Even a fully "
    "compromised server administrator cannot forge a signature, because the private keys never leave the "
    "signers’ TPMs and the server only ever sees public keys."
)

h2("3.7  Supporting Features")
para(
    "Beyond the three pillars, the system provides several supporting features that round out the review "
    "experience. Passage-level highlights let reviewers drag-select a region of the PDF and attach a "
    "comment anchored to the exact passage; highlights are stored in document_highlights with "
    "page-relative bounding rectangles, are version-scoped, and follow the same access rules as approvals "
    "(a reviewer may highlight only a document they hold an approval row for in the current round). "
    "Review deadlines and reminders let the owner set a deadline that is stamped on every approval row in "
    "the round; the dashboard and review queue colour-code pending rows (overdue in red, due within "
    "twenty-four hours in amber) and a lazy-reminder pattern inserts a reminder notification whenever a "
    "reviewer loads the dashboard or review queue and an item is overdue, avoiding the need for a "
    "separate cron job."
)
para(
    "The dashboard surfaces three lightweight analytics — Documents by Status, Documents by Month, and "
    "the Approval / Rejection Ratio — each rendered as a small server component running a single "
    "aggregate SQL query and emitting inline-styled HTML, so no charting library or extra client "
    "JavaScript is shipped. Role-aware quick-action cards (New Document, Manage Documents, Awaiting "
    "Review, Needs Revision, and, for reviewers and administrators, Review Queue and Admin Panel) deep-"
    "link into filtered document views. Realtime live updates are provided by Supabase Realtime: a "
    "NotificationBell component and an invisible DashboardRealtime component subscribe to the relevant "
    "tables and trigger a debounced refresh so the UI reflects new state within a second of any change, "
    "without the user reloading. Notably, realtime events are gated by the same RLS policies that govern "
    "SELECT queries, so a malicious client cannot subscribe to changes it has no right to see. Finally, "
    "every event that produces an in-app notification also dispatches an email through Brevo (eight email "
    "types, from reviewer assignment to account approval); sends are best-effort and wrapped in a "
    "try-catch, so a Brevo outage cannot block the rest of a request."
)

h2("3.8  AI Rate Limiting")
para(
    "The two summarization and question-answering endpoints, together with the text-extraction endpoint, "
    "enforce a per-user rate limit to bound how frequently a single account can invoke the paid AI "
    "features. Each user is allowed a configurable number of AI calls per time window — twenty calls per "
    "ten minutes by default, shared across the three endpoints — after which the route returns HTTP 429 "
    "with a Retry-After header and a message indicating how long to wait. The counter is maintained in "
    "the PostgreSQL database rather than in application memory, because the application runs on stateless "
    "serverless functions where an in-process counter would reset on every cold start and would not be "
    "shared between concurrent instances. A single SECURITY DEFINER database function performs the "
    "check-and-increment atomically, so two concurrent requests cannot both pass the limit. The limit "
    "fails open: if the counter query itself fails, the request is allowed rather than blocked, so a "
    "fault in the rate-limiting layer cannot disable the AI features."
)

h2("3.9  REST API Surface")
para(
    "Table 3.5 lists the main API endpoints. All use JSON for request and response bodies, return "
    "appropriate HTTP status codes for errors, and require an authenticated cookie session; endpoints "
    "performing privileged writes additionally verify the caller’s role."
)
api_headers = ["Method", "Endpoint", "Description"]
api_rows = [
    ["POST", "/api/auth/register-notify", "Fire-and-forget; notifies every approved admin (in-app + "
     "email) that a new user awaits approval."],
    ["POST", "/api/profile/webauthn/register-options", "Returns a WebAuthn registration challenge for "
     "first-time signing-key setup."],
    ["POST", "/api/profile/webauthn/register", "Verifies the attestation; stores the user’s public key, "
     "AAGUID, and counter."],
    ["POST", "/api/documents", "Create a new document and register its first version after server-side "
     "validation."],
    ["POST", "/api/documents/[id]/versions", "Register a new version of an existing document after "
     "validation."],
    ["GET", "/api/documents/[id]/file-hash", "Returns the cached SHA-256 hash of the latest version (the "
     "WebAuthn challenge)."],
    ["POST", "/api/documents/[id]/submit", "Submit for review with reviewer list, deadline, and the "
     "owner’s WebAuthn signature."],
    ["POST", "/api/approvals/[approvalId]/decide", "Record an approve (with WebAuthn signature) or reject "
     "(comment only) decision."],
    ["POST", "/api/documents/[id]/extract-text", "Run the hybrid PDF text-extraction pipeline for the "
     "latest version."],
    ["POST", "/api/documents/[id]/ai-summary", "Run extraction and generate a structured AI summary for "
     "the latest version."],
    ["POST", "/api/documents/[id]/ai-chat", "Answer a natural-language question about the document via the "
     "OpenAI API."],
    ["POST", "/api/documents/[id]/highlights", "Add a passage-level highlight with a comment to the "
     "current version."],
    ["DELETE", "/api/documents/[id]/highlights/[highlightId]", "Delete a passage highlight (author or "
     "administrator only)."],
    ["GET", "/api/documents/[id]/signature/verify", "Verify every stored signature against the current "
     "file and each signer’s public key."],
    ["PATCH", "/api/admin/users/[id]/status", "Administrator: approve or reject a pending account."],
    ["PATCH", "/api/admin/users/[id]/role", "Administrator: change a user’s role (employee / reviewer / "
     "admin)."],
    ["DELETE", "/api/admin/documents/[id]", "Administrator: delete a document with all dependent rows and "
     "storage files."],
]
make_table(api_headers, api_rows, widths=[0.7, 2.5, 2.8], body_size=9)
caption("Table 3.5: Main REST API endpoints exposed by the backend.")

h2("3.10  Testing and Validation")
para(
    "Testing focuses on functional correctness of the workflow and the effectiveness of the security "
    "controls. Each main use case is exercised manually with multiple test accounts in different roles: "
    "an employee uploads a document and submits it to two reviewer accounts, observing notifications and "
    "dashboard updates; the reviewers approve or reject; and the resulting transitions are checked "
    "against the expected behavior of the state machine. Edge cases such as resubmission after rejection, "
    "late approvals, and verification of a tampered file are exercised explicitly."
)
para(
    "Row-Level Security policies are tested through two complementary approaches: using the application "
    "interface as different users to confirm each sees only the documents and approvals they are "
    "entitled to, and using the Supabase SQL editor with role impersonation (setting request.jwt.claims "
    "inside a transaction) to exercise the policies directly without any application code. API endpoints "
    "are exercised both through the UI and through direct HTTP requests under three scenarios — an "
    "authorized user with a valid payload, an unauthorized user, and an authorized user with a malformed "
    "payload — with expected outcomes of success, 403 Forbidden, and 400 Bad Request respectively. The "
    "manual testing confirms correct behavior across the main use cases; systematic automated testing "
    "with a framework such as Playwright is identified in Chapter 5 as future work."
)
page_break()

# =====================================================================
# CHAPTER 4 - RESULTS & DISCUSSION   (target ~6 pages)
# =====================================================================
h1("Chapter 4  Results and Discussion")

h2("4.1  Functional Coverage")
para(
    "The implemented prototype covers all functional and non-functional requirements. Employees can "
    "register accounts, wait for administrator approval, sign in, upload PDFs, submit them for "
    "multi-reviewer review with deadlines, respond to rejections with new versions, and consult the AI "
    "assistant. Reviewers can list documents assigned to them, approve or reject with comments, and leave "
    "passage-level highlights. Administrators can manage user accounts and roles, browse all documents "
    "and approvals, search by extracted text, and inspect the audit log. Table 4.1 maps each functional "
    "requirement to its implementation status; all were exercised during manual testing."
)
fr_headers = ["ID", "Requirement", "Status"]
fr_rows = [
    ["FR1", "Account registration with administrator approval", "Implemented"],
    ["FR2", "Authenticated sign-in and sign-out via cookies", "Implemented"],
    ["FR3", "Document creation and PDF upload", "Implemented"],
    ["FR4", "Multi-version uploads with version history", "Implemented"],
    ["FR5", "Submission with reviewer selection and deadline", "Implemented"],
    ["FR6", "Approval or rejection with optional or required comment", "Implemented"],
    ["FR7", "Unanimous approval and early-termination rejection", "Implemented"],
    ["FR8", "Passage-level highlights and comments on PDFs", "Implemented"],
    ["FR9", "AI summary, key points, risk notes, and Q&A", "Implemented"],
    ["FR10", "Multi-party WebAuthn signing and Verify Integrity covering every signature", "Implemented"],
    ["FR11", "Printable approval certificate page", "Implemented"],
    ["FR12", "In-app notifications and lazy reminders", "Implemented"],
    ["FR13", "Administrator user and role management", "Implemented"],
    ["FR14", "Administrator document oversight with full-text search", "Implemented"],
    ["FR15", "Audit log browsing with filters", "Implemented"],
    ["FR16", "Realtime updates for notifications and dashboard", "Implemented"],
    ["FR17", "Transactional email notifications via Brevo", "Implemented"],
]
make_table(fr_headers, fr_rows, widths=[0.7, 4.6, 1.2], body_size=10)
caption("Table 4.1: Coverage of functional requirements in the implemented prototype.")

h2("4.2  Performance Evaluation")
para(
    "Performance was evaluated by measuring representative interactions against a production build (next "
    "build followed by next start) running on a laptop with an AMD Ryzen 7 6800H processor and "
    "thirty-two gigabytes of memory, with the Supabase backend hosted in the Singapore region. Each "
    "figure is the average of five consecutive measurements taken from the browser’s network timing "
    "panel under light load. Page navigation typical of a logged-in user completes within one second, "
    "satisfying NFR1, and generating a structured AI summary for a representative ten-page document takes "
    "between three and seven seconds, with most of the time spent in the OpenAI API call rather than the "
    "server, well within NFR2. Table 4.2 reports the measured times."
)
perf_headers = ["Operation", "Avg time", "Notes"]
perf_rows = [
    ["Dashboard load", "~600 ms", "After query parallelization"],
    ["Document detail", "~700 ms", "Including PDF render"],
    ["AI summary (10-page)", "4.2 s", "Most time in OpenAI call"],
    ["AI Q&A", "2.8 s", ""],
    ["Verify Integrity", "1.1 s", "Hash recomputation + per-signature crypto re-verification"],
    ["WebAuthn signature verify", "~150 ms", "Server-side verifyAuthenticationResponse"],
]
make_table(perf_headers, perf_rows, widths=[2.0, 1.2, 2.8], body_size=10)
caption("Table 4.2: Measured response times for representative interactions under light load.")

h2("4.3  Security Evaluation")
para(
    "Security was evaluated by attempting to bypass the access controls in several ways. As an employee, "
    "an attempt to read another user’s documents through the UI correctly returned a not-found page; the "
    "same attempt by sending a direct request with a manually crafted document identifier was rejected by "
    "Row-Level Security at the database level even when the application-level check was deliberately "
    "removed in a test branch. An attempt to upload a non-PDF file with a fake .pdf extension was caught "
    "by the magic-byte validation, and the staging object was deleted as expected. Table 4.3 summarizes "
    "the attack scenarios tested and the layer that blocked each, confirming that the layered controls "
    "work in practice."
)
sec_headers = ["Attack scenario", "Method", "Result"]
sec_rows = [
    ["Cross-user document read", "Direct URL access", "Blocked (RLS)"],
    ["Bypass app-level role check", "Removed requireRole guard", "Blocked (RLS)"],
    ["Non-PDF upload", "Renamed .exe to .pdf", "Caught (magic-byte check)"],
    ["Oversized upload", "50 MB PDF", "Rejected (size limit)"],
    ["Tamper after approval", "Replaced file in Storage", "Detected (approved_hash mismatch)"],
    ["SQL injection (search)", "'; DROP TABLE --", "Safe (parameterized queries)"],
]
make_table(sec_headers, sec_rows, widths=[2.0, 2.0, 2.0], body_size=10)
caption("Table 4.3: Security attack scenarios tested against the implemented prototype.")

h2("4.4  Discussion of Research Questions")
para(
    "RQ1 (modeling multi-role workflows) is addressed positively. The state machine is expressive enough "
    "to capture single- and multi-reviewer scenarios, resubmission after rejection, and the integration "
    "of signing, while remaining simple enough to implement in a small set of API routes and a few SQL "
    "constraints. The use of round numbers in the approvals table provides a clean way to preserve "
    "history while focusing the live progress display on the current round.",
)
para(
    "RQ2 (integrating large language models) is also addressed positively, with caveats. Structured "
    "outputs in the form of JSON-schema-constrained responses proved the single most valuable technique "
    "for making the summary feature reliable, since the application parses the response without fragile "
    "string matching. The hybrid extraction pipeline handles both text-based and scanned PDFs within the "
    "page limit, and the cost-control measures (input truncation, question caps, logged usage, and the "
    "per-user rate limit) keep cost predictable. The caveats are that the system depends on an external "
    "paid service and that quality depends on the underlying model, which is outside the project’s "
    "control."
)
para(
    "RQ3 (layered security) is supported by the tests above. The combination of page-level checks via "
    "requireUser and requireRole, server-mediated writes through the service-role client after explicit "
    "permission checks, and SELECT-only RLS on the data tables provides a defense-in-depth strategy "
    "robust to several classes of bugs. The deliberate decision to perform all inserts into documents and "
    "document_versions through the server, rather than directly from the anon-key client, makes it "
    "impossible for a malicious client to bypass the upload validation pipeline."
)

h2("4.5  User Interface Walkthrough")
para(
    "This section illustrates the implemented user interface. Figure 4.1 shows the sign-in and "
    "registration screen, the entry point of the application. Figure 4.2 shows the employee dashboard, "
    "combining metric cards, the Documents by Status and Documents by Month analytics charts, and "
    "role-aware quick-action cards that link directly to filtered document views; the document list "
    "itself supports status filtering with live counts. Figure 4.3 shows the document detail page — the "
    "main workspace — combining an inline PDF viewer with passage highlights, the upload and submission "
    "controls, and the AI workspace. Figure 4.4 zooms in on the AI workspace after the summary has been "
    "generated, showing the three structured fields, and Figure 4.5 shows the printable approval "
    "certificate with live verification status."
)
figure_placeholder("Figure 4.1", "Login and registration interface")
caption("Figure 4.1: Login and registration interface.")
figure_placeholder("Figure 4.2", "Employee dashboard with document overview")
caption("Figure 4.2: Employee dashboard with document overview.")
figure_placeholder("Figure 4.3", "Document detail page with inline PDF viewer")
caption("Figure 4.3: Document detail page with inline PDF viewer.")
figure_placeholder("Figure 4.4", "AI assistant panel showing the generated summary")
caption("Figure 4.4: AI assistant panel showing the generated summary.")
figure_placeholder("Figure 4.5", "Approval certificate and integrity verification page")
caption("Figure 4.5: Approval certificate and integrity verification page.")

h2("4.6  Limitations Observed in Practice")
para(
    "Several limitations became apparent during testing. First, the OCR fallback is restricted to "
    "documents of up to ten pages; longer scanned documents return an error and must be split manually. "
    "This is a deliberate cost-control decision that could be relaxed with a dedicated OCR service. "
    "Second, WebAuthn signing requires a platform authenticator (Windows Hello, Touch ID, or mobile "
    "biometric); on Linux desktops without one, users cannot sign with the current strict configuration. "
    "The system detects this and shows a setup-instructions panel; allowing cross-device WebAuthn (phone "
    "via QR code) is a one-line change but would weaken the “credential bound to this physical machine” "
    "property and was therefore not enabled by default. Third, emails are sent from a free Gmail address "
    "that is not domain-authenticated, so Gmail routes the first message to the Spam folder until the "
    "recipient marks it as not spam; authenticating a custom domain in Brevo would remove this. Fourth, "
    "no end-to-end automated test suite has been written, so regressions can in principle slip through. "
    "Fifth, the AI assistant relies on a third-party service, so document content leaves the deployment "
    "boundary when summarization or question-answering is invoked; organizations with regulated data "
    "would need to disable AI for such documents or migrate to a self-hosted model. Sixth, text-layer "
    "extraction quality varies by document type: word-processor and LaTeX article PDFs extract cleanly, "
    "while slide decks and documents with heavy mathematical notation can run word boundaries together, "
    "and OCR degrades on complex multi-column layouts such as newspaper or magazine pages. This is "
    "consistent with the project’s intended scope of single-column organizational documents; a more "
    "sophisticated multi-engine extraction pipeline would address it."
)
page_break()

# =====================================================================
# CHAPTER 5 - CONCLUSION & FUTURE WORK   (target ~3 pages)
# =====================================================================
h1("Chapter 5  Conclusion and Future Work")

h2("5.1  Conclusion")
para(
    "This report has presented the design, implementation, and evaluation of an AI-assisted document "
    "approval system that combines a multi-role review workflow, integrity verification, and an "
    "integrated AI assistant. The system is implemented as a single Next.js application backed by "
    "Supabase for data and storage and by the OpenAI API for language model capabilities, allowing a "
    "single developer to deliver, within the scope of a graduation project, a system that would "
    "traditionally require the integration of several separate products."
)
para(
    "Addressing RQ1, the workflow has been modeled as a finite state machine over documents and a "
    "round-based design over approvals — expressive enough to capture multi-reviewer scenarios with "
    "resubmission and simple enough to implement with a small number of API routes and SQL constraints, "
    "and enforced consistently across the application, the database, and the user interface. Addressing "
    "RQ2, the integration of the OpenAI API provides genuine value to reviewers, with structured outputs "
    "as the key technique that makes the summary feature reliable, a hybrid extraction pipeline handling "
    "both digital and scanned PDFs within a defined cost envelope, and graceful degradation through "
    "clear status codes when the model is unavailable or rate-limited. Addressing RQ3, the layered "
    "security architecture combining application-level guards with Row-Level Security prevents "
    "unauthorized access even in the presence of bugs or compromised clients, and the staging-then-"
    "validate upload pipeline ensures no invalid file can ever appear in the storage area or the "
    "database. Beyond these questions, the project demonstrates that a modern web stack, combined with "
    "managed services and a carefully scoped feature set, allows a single student to build a non-trivial "
    "production-quality application within a graduation-project timeframe."
)

h2("5.2  Future Work")
para("Several directions for future work were identified during implementation and evaluation, grouped "
     "by priority and expected impact.")

h3("5.2.1  High Priority")
bullet("Embed signatures into the PDF binary. The current implementation stores signatures in a separate "
       "table; embedding the WebAuthn signature into the PDF’s native signature dictionary (PAdES / "
       "PKCS#7) would allow standalone verification in Adobe Reader without this application’s database. "
       "This requires no cryptographic changes — the same ECDSA bytes can be wrapped in the PAdES "
       "container — but adds PDF-manipulation complexity (pdf-lib or @signpdf/signpdf).")
bullet("Trusted timestamp authority (TSA) integration. The signed_at timestamp is currently recorded by "
       "the application server and trusted only as much as the database is. For non-repudiation that "
       "survives database tampering, the application should request a cryptographic timestamp from an "
       "external TSA (per RFC 3161) at signing time and store the TSA signature as part of the evidence, "
       "as real-world e-signature platforms and Vietnamese providers do for compliance with the Law on "
       "Electronic Transactions.")

h3("5.2.2  Medium Priority")
bullet("Automated end-to-end tests with a Playwright-based suite covering the main user flows.")
bullet("Extraction-quality heuristic that also routes structurally-broken text-layer extractions (for "
       "example slide decks) through the vision model, improving summary quality on heavily-formatted "
       "documents.")
bullet("Dedicated OCR service (Tesseract or a cloud OCR API) to process larger scanned documents without "
       "raising costs.")
bullet("Self-hosted open-weights model (such as Llama 3) to eliminate per-call cost and the need to "
       "transmit document content to a third party, addressing the data-sovereignty limitation; the "
       "trade-off is GPU infrastructure and the likely need for a 30B+ model to match gpt-5.4-mini "
       "quality.")
bullet("Mobile-responsive redesign, document templates and tagging, and multi-device / cross-device "
       "WebAuthn credentials (including a configurable per-deployment toggle for phone-based signing).")

h3("5.2.3  Lower Priority and Operational Improvements")
bullet("Research directions: multi-language summaries, a risk-classification model trained on existing "
       "risk notes, and configurable workflow definitions (for example “two of three reviewers must "
       "approve” or staged sequential review).")
bullet("Operational: production monitoring (structured logging, error tracking via a service such as "
       "Sentry), machine-readable OpenAPI documentation, backup and disaster-recovery procedures, email "
       "domain authentication (SPF / DKIM / DMARC in Brevo), and WebAuthn credential lifecycle "
       "management (revoke, rotate, and re-enroll flows).")
page_break()

# =====================================================================
# BIBLIOGRAPHY  (kept; not part of the 35-page budget)
# =====================================================================
h1("Bibliography")
refs = [
    "Vercel, Next.js Documentation, https://nextjs.org/docs, accessed 2026.",
    "Meta Open Source, React Documentation, https://react.dev, accessed 2026.",
    "Tailwind Labs, Tailwind CSS Documentation, https://tailwindcss.com/docs, accessed 2026.",
    "Supabase Inc., Supabase Documentation, https://supabase.com/docs, accessed 2026.",
    "PostgreSQL Global Development Group, PostgreSQL Documentation: Row Security Policies, "
    "https://www.postgresql.org/docs/current/ddl-rowsecurity.html, accessed 2026.",
    "OpenAI, OpenAI API Reference, https://platform.openai.com/docs/api-reference, accessed 2026.",
    "OpenAI, Structured Outputs Guide, https://platform.openai.com/docs/guides/structured-outputs, "
    "accessed 2026.",
    "Mozilla Developer Network, PDF.js Project, https://mozilla.github.io/pdf.js/, accessed 2026.",
    "World Wide Web Consortium, Web Cryptography API, W3C Recommendation, January 2017, "
    "https://www.w3.org/TR/WebCryptoAPI/.",
    "National Institute of Standards and Technology, Secure Hash Standard (SHS), FIPS Publication 180-4, "
    "August 2015.",
    "Internet Engineering Task Force, JSON Web Token (JWT), RFC 7519, May 2015.",
    "D. F. Ferraiolo and D. R. Kuhn, Role-Based Access Control, in Proceedings of the 15th National "
    "Computer Security Conference, Baltimore, MD, October 1992, pp. 554-563.",
    "R. S. Sandhu, E. J. Coyne, H. L. Feinstein and C. E. Youman, Role-Based Access Control Models, "
    "IEEE Computer, vol. 29, no. 2, pp. 38-47, February 1996.",
    "L. Richardson and S. Ruby, RESTful Web Services, O’Reilly Media, 2007.",
    "L. Bass, P. Clements and R. Kazman, Software Architecture in Practice, 4th ed., Addison-Wesley, "
    "2021.",
    "Object Management Group, Unified Modeling Language (UML) Version 2.5.1, OMG Document "
    "formal/2017-12-05, December 2017.",
    "T. H. Cormen, C. E. Leiserson, R. L. Rivest and C. Stein, Introduction to Algorithms, 3rd ed., "
    "MIT Press, 2009.",
    "M. Stonebraker and L. A. Rowe, The Design of POSTGRES, ACM SIGMOD Record, vol. 15, no. 2, "
    "pp. 340-355, 1986.",
    "World Wide Web Consortium, Web Authentication: An API for accessing Public Key Credentials "
    "(WebAuthn) Level 3, W3C Recommendation, October 2023, https://www.w3.org/TR/webauthn-3/.",
    "FIDO Alliance, FIDO2: WebAuthn & CTAP, https://fidoalliance.org/fido2/, accessed 2026.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Inches(0)
    p.add_run(f"[{i}] {r}").font.size = Pt(12)

doc.save(OUT)
print("SAVED:", OUT)
print("Paragraphs:", len(doc.paragraphs))
print("Tables:", len(doc.tables))
